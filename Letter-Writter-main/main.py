# This program takes a list of property information and generates a personalized letter based on the
# specified information
# Written by William Andrew Spicher 12/21/2021
import tkinter.messagebox

import docx.shared
from docx import Document as Document
from docxcompose.composer import Composer
from datetime import date
import tkinter as Tk
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import askdirectory
from tkinter.messagebox import *
from tkinter import ttk
import os

# Global Variables
globalInputList = ""
globalOutputDir = ""
globalTaxSaleDate = ""

# The property class contains all necessary information necessary to compose the output letter
class Property:
    # Constructor
    def __init__(self, propAddress, ownerName, ownerAddress, ownerType):
        self.propAddress = propAddress
        self.ownerName = ownerName
        self.ownerAddress = ownerAddress
        self.ownerType = ownerType

# This function takes the property input list and stores the items into a list of tuples
def propertyListGen(filePath):

    outputList = []
    tempListIndex = 0
    tempString = ""
    tempList = ["","","",""]

    # Open the file
    listFile = open(filePath, "r")

    # Iterate through list
    while True:
        line = listFile.readline().strip("\n\t")

        # Check to see of the end of file has been reached
        if line != '':
            # Check to see if the line is a block end
            if line != "====================================================":
                # Check to see if the line is a separator
                if line != "---":
                    # Store information
                    tempString += line + " "
                else: # if tempString equals "---"
                    # Add temp string to appropriate spot in the tempList
                    tempList[tempListIndex] = tempString.strip(" ")

                    # Reset tempString and increment index
                    tempString = ""
                    tempListIndex += 1
            else:
                # Add temp string to appropriate spot in the tempList
                tempList[tempListIndex] = tempString.strip(" ")

                # Create a Property object with the info and append it to the output list
                tempProperty = Property(tempList[0], tempList[1], tempList[2], tempList[3])
                outputList.append(tempProperty)

                # Reset tempString, index, and tempList
                tempString = ""
                tempListIndex = 0
                for i in range(4):
                    tempList[i] = ""

        else:
            # EOF reached
            listFile.close()
            return outputList

# This function takes the current date and returns it into a longer form
def dateExpander(currentDate):
    # year- month- day
    outDate = ""
    monthDict = {
        "1" : "January",
        "2" : "February",
        "3" : "March",
        "4" : "April",
        "5" : "May",
        "6" : "June",
        "7" : "July",
        "8" : "August",
        "9" : "September",
        "10": "October",
        "11": "November",
        "12": "December"
    }
    outDate += monthDict[str(currentDate.month)] + " " + str(currentDate.day) +  ", " +str(currentDate.year)
    return outDate

# This function takes a list of Property objects and generates a personalized letter with
# with the objects' specific information
def letterGenerator(propertyList, saleDate, tempFilePath, outputFilePath):
    documentList = []
    currentDate = dateExpander(date.today())

    # Loop through all properties in Properties list
    for property in propertyList:
        # Open new template document copy
        doc = Document(tempFilePath)

        # Change document values to Property object values
        for paragraph in doc.paragraphs:
            # Current date
            paragraph.text = paragraph.text.replace("currentDate", currentDate)

            # Make sure greeting is correct
            if property.ownerType == "company":
                # Property owned by company
                paragraph.text = paragraph.text.replace("ownerFirstName", "Sir or Madam")
            elif property.ownerType == "person":
                # Property owned by a person, format first name
                formatOwnerName = str(property.ownerName.split()[0].lower().capitalize())

                # Add formatted first name to the letter opener
                paragraph.text = paragraph.text.replace("ownerFirstName", formatOwnerName)
            else:
                # Error, there is a typo in the owner type on the input list for this property
                print(f"ERROR: Typo in owner type on input list for item: {property.ownerName}")
                paragraph.text = paragraph.text.replace("ownerFirstName", "Sir or Madam")

            # Property Address
            paragraph.text = paragraph.text.replace("propertyAddress", str(property.propAddress) + ", GA")

            # Tax Sale Date
            paragraph.text = paragraph.text.replace("taxSaleDate", saleDate)

            # Bottom page content
            paragraph.text = paragraph.text.replace("______________________________________________________________________________",
                                                    "_____________________________________________________________________________________")

        # Change owner name and address at top left of template letter
        # Find table in template
        table = doc.tables[0]

        # Adjust table rows and column to ensure format is maintained
        # table.columns[0].width = docx.shared.Inches(2.4)
        for cell in table.columns[0].cells:
            cell.width = docx.shared.Inches(2.3)

        # Assign Text
        table.rows[0].cells[0].text = property.ownerName + "\n" + property.ownerAddress

        # Change the greeting font to match the example letter
        # Chaparral Pro
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Chaparral Pro"
        # Save as new document
        tempDocName = outputFilePath + property.ownerName + "_Letter.docx"
        doc.save(tempDocName)

        # Add document to the document list
        documentList.append(tempDocName)
    return documentList

# This function sets the global filepath of the input list
def setInputList(statusLabel):
    global globalInputList
    globalInputList = askopenfilename(title="Select Input List", multiple=False)

    # Update status label
    statusLabel.config(text=globalInputList)

# This function sets the global filepath of the output directory
def setOutputDir(statusLabel):
    global globalOutputDir
    globalOutputDir = askdirectory(title="Select Output Folder")

    # Update status label
    statusLabel.config(text=globalOutputDir)

# This function sets the global tax sale date
def setTaxSaleDate(month, day, year, statusLabel):
    global globalTaxSaleDate
    globalTaxSaleDate = f"{month} {day}, {year}"

    # Update status label
    statusLabel.config(text=f"Confirm Date: {globalTaxSaleDate}")

# This funtion combines a list of docx documents into one document
def documentCombiner(masterFileName, docList):
    # Create a master document
    mergedDoc = Document()

    for index, file in enumerate(docList):
        tempDoc = Document(file)
        # dont add page break if last file is reached
        if index < len(docList)-1:
            tempDoc.add_page_break()

        # Loop through each body and append contents
        for element in tempDoc.element.body:
            mergedDoc.element.body.append(element)

    # Save the combined file
    mergedDoc.save(masterFileName)

    # Delete all files in the docList from directory
    for file in docList:
        os.remove(file)

# This function performs all steps of letter generation upon the pressing of the run button
def startEntireProg():
    # Perform checks to make sure all necessary data has been gathered
    if globalInputList == "":
        # Spit out error
        Tk.messagebox.showerror("Missing Input List", "Please select an input list.")
        # Return to not proceed further
        return
    elif globalOutputDir == "":
        Tk.messagebox.showerror("Missing Output Folder", "Please select an output folder.")
        return
    elif globalTaxSaleDate == "":
        Tk.messagebox.showerror("Missing Tax Sale Date", "Please select a tax sale date.")
    else:
        # No input errors
        # Generate property list
        proplist = propertyListGen(globalInputList)

        # Generate letters from property list
        letterList = letterGenerator(proplist, globalTaxSaleDate, "template.docx", "")

        # Combine documents into single document
        documentCombiner(f"{globalOutputDir}/{globalTaxSaleDate.split()[0]}_letters.docx", letterList)

        # Tell user process is complete
        Tk.messagebox.showinfo("Letters Generated", "Property letters have been generated.")





# This function instantiates the GUI for the program
def progGUI():
    # Colors
    lightBlue = "#B7D8D6"
    darkerBlue = "#789E9E"
    darkestBlue = "#4D6466"
    tan = "#EEF3D8"
    lightRed = "#FE615A"

    # Declare a window object
    window = Tk.Tk()

    # Set window title
    window.title("Tax Sale Letter Writter")

    # Set window dimensions
    #window.configure(width=500, height=500)
    window.geometry('700x500')
    window.grid_rowconfigure(1, weight=1)
    window.grid_columnconfigure(0, weight=1)

    # Set window background color
    window.configure(bg=darkestBlue)

    # Top label containing the program name
    topLabelName = Tk.Label(window, text="Tax Sale Letter Writer", bg = darkestBlue, fg = tan, font=("Calibri", 25)).pack(side='top')
    topLabelDev = Tk.Label(window, text="by Andrew Spicher", bg = darkestBlue, fg = tan, font=("Calibri", 15)).pack(side='top')

    # Define a frame for widget organization
    mainFrame = Tk.Frame(window)
    mainFrame.configure(bg=darkestBlue)

    # Input list selector configurations
    listSelectorLabel = Tk.Label(mainFrame, text="Select Input File:", font=("Calibri", 20), bg=darkestBlue, fg=tan)
    listPathLabel = Tk.Label(mainFrame, text="", font=("Calibri", 14),bg = darkestBlue, fg = lightBlue)
    listSelectorButton = Tk.Button(mainFrame, text="Browse", font=("Calibri", 15), bg=darkestBlue, command = lambda: setInputList(listPathLabel))
    listSelectorLabel.grid(column=0, row=1, sticky = 'W')
    listSelectorButton.grid(column=1, row=1)
    listPathLabel.grid(column=0, row=3,sticky = 'W')

    # Separator label
    sepLab = Tk.Label(mainFrame, text="\n", bg = darkestBlue).grid(column=0, row=4)

    # Output directory selector configurations
    outputSelectorLabel = Tk.Label(mainFrame, text="Select Output Folder:", font=("Calibri", 20), bg = darkestBlue, fg = tan)
    outputPathLabel = Tk.Label(mainFrame, text="", font=("Calibri", 14),bg = darkestBlue, fg = lightBlue)
    outputSelectorButton = Tk.Button(mainFrame, text="Browse", font=("Calibri", 15), bg=darkestBlue, command = lambda: setOutputDir(outputPathLabel))

    outputSelectorLabel.grid(column=0, row=5, sticky = 'W')
    outputSelectorButton.grid(column=1, row=5)
    outputPathLabel.grid(column=0,row=6, sticky='W')

    # Separator label
    sepLab2 = Tk.Label(mainFrame, text="\n", bg = darkestBlue).grid(column=0, row= 7)

    # Tax sale date configurations
    dateLabel = Tk.Label(mainFrame, text="Enter Tax Sale Date:",font=("Calibri", 20), bg = darkestBlue, fg = tan)
    dateLabel.grid(column=0, row=8, sticky='W')

    # Month
    n = Tk.StringVar
    tsMonth = ttk.Combobox(mainFrame, width=10, textvariable= n, background = darkestBlue)
    tsMonth ['values'] = ('January',
                          'February',
                          'March',
                          'April',
                          'May',
                          'June',
                          'July',
                          'August',
                          'September',
                          'October',
                          'November',
                          'December')
    # Pack month
    tsMonth.grid(column = 0, row = 9, sticky='W', padx = 0)
    # Set default month value
    tsMonth.current(0)

    # Day
    tsDay = ttk.Combobox(mainFrame, width=5, textvariable = n, background = darkestBlue)
    tsDay["values"] = ('1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13',
                       '14', '15', '16', '17', '18', '19', '20', '21', '22', '23', '24', '25', '26',
                       '27', '28', '29', '30', '31')
    # Pack day
    tsDay.grid(column = 1, row = 9, sticky='W', padx = 0)
    # Set default day value
    tsDay.current(0)

    # Year
    tsYear = ttk.Combobox(mainFrame, width=10, textvariable = n, background = darkestBlue)
    tsYear["values"] = ('2022', '2023', '2024', '2025', '2026','2027', '2028','2029','2030')
    # Pack year
    tsYear.grid(column = 2, row = 9, sticky='W', padx = 0)
    # Set default year value
    tsYear.current(0)

    # Date confirmation
    dateConfirmLabel = Tk.Label(mainFrame, text="", font=("Calibri", 14), bg=darkestBlue, fg=lightBlue)
    dateConfirmButton = Tk.Button(mainFrame, text="Confirm", font=("Calibri", 15), bg=darkestBlue,
                                  command = lambda: setTaxSaleDate(tsMonth.get(), tsDay.get(), tsYear.get(), dateConfirmLabel))
    # Pack date confirm button and label
    dateConfirmButton.grid(column=0, row=10, sticky='W')
    dateConfirmLabel.grid(column=1, row=10, sticky='W')

    # Separator label
    sepLab3 = Tk.Label(mainFrame, text="\n", bg=darkestBlue).grid(column=0, row=11)
    # Run button
    runButton = Tk.Button(mainFrame, text="Run",font=("Calibri", 15), bg=darkestBlue, width= 10,
                          command = startEntireProg)
    runButton.grid(column=2, row=12, sticky='E')

    # Pack the main frame into the window
    mainFrame.pack(side='left')

    # Run window
    window.mainloop()

# TODO:
# Fix document formatting error at bottom
# Test program

# Test GUI
progGUI()





